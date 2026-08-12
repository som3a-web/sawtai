"""Secure file extraction, object storage, and embedding adapters."""

import asyncio
import hashlib
import io
import math
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlsplit

import httpx
import pypdfium2 as pdfium
import pytesseract
from docx import Document
from minio import Minio
from pypdf import PdfReader

from sawtai.arabic import normalize_for_search
from sawtai.config import Settings

EMBEDDING_DIMENSION = 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
EICAR_MARKER = b"EICAR-STANDARD-ANTIVIRUS-TEST-FILE"
PDF_ACTIVE_MARKERS = (b"/JavaScript", b"/JS ", b"/Launch", b"/EmbeddedFile")


class DocumentIngestionError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


@dataclass(frozen=True)
class ValidatedUpload:
    filename: str
    extension: str
    media_type: str
    sha256: bytes


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    method: str
    page_count: int | None
    warnings: tuple[str, ...]


@dataclass(frozen=True)
class EmbeddingBatch:
    vectors: tuple[tuple[float, ...], ...]
    provider: str


def sanitize_filename(filename: str) -> str:
    clean = Path(filename).name.strip().replace("\x00", "")
    clean = re.sub(r"[^\w.\-\u0600-\u06ff ]+", "_", clean, flags=re.UNICODE)
    if not clean:
        raise DocumentIngestionError("invalid_filename", "The uploaded filename is invalid")
    return clean[:180]


def validate_upload(
    *,
    filename: str,
    content: bytes,
    declared_media_type: str | None,
    max_bytes: int,
) -> ValidatedUpload:
    clean_name = sanitize_filename(filename)
    extension = Path(clean_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentIngestionError("unsupported_type", "Only PDF, DOCX, TXT, and Markdown files are accepted")
    if not content:
        raise DocumentIngestionError("empty_file", "The uploaded file is empty")
    if len(content) > max_bytes:
        raise DocumentIngestionError("file_too_large", f"The file exceeds the {max_bytes // 1_048_576} MB limit")
    if EICAR_MARKER in content:
        raise DocumentIngestionError("malware_detected", "The uploaded file failed malware screening")
    media_type = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }[extension]
    if extension == ".pdf":
        if not content.startswith(b"%PDF-"):
            raise DocumentIngestionError("signature_mismatch", "The file extension does not match PDF content")
        if any(marker in content for marker in PDF_ACTIVE_MARKERS):
            raise DocumentIngestionError("active_content", "PDF files containing active or embedded content are not accepted")
    elif extension == ".docx":
        _validate_docx_archive(content)
    else:
        if b"\x00" in content[:4096]:
            raise DocumentIngestionError("signature_mismatch", "The text file contains binary content")
    if declared_media_type and declared_media_type not in {media_type, "application/octet-stream"}:
        raise DocumentIngestionError("mime_mismatch", "The declared file type does not match its extension")
    return ValidatedUpload(
        filename=clean_name,
        extension=extension,
        media_type=media_type,
        sha256=hashlib.sha256(content).digest(),
    )


def _validate_docx_archive(content: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            members = archive.infolist()
            names = {member.filename for member in members}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise DocumentIngestionError("signature_mismatch", "The file is not a valid DOCX document")
            if len(members) > 5_000:
                raise DocumentIngestionError("archive_limit", "The DOCX archive contains too many entries")
            total_size = sum(member.file_size for member in members)
            compressed_size = max(1, sum(member.compress_size for member in members))
            if total_size > 100 * 1_048_576 or total_size / compressed_size > 120:
                raise DocumentIngestionError("archive_limit", "The DOCX archive exceeds safe expansion limits")
            if any("vbaproject.bin" in name.lower() for name in names):
                raise DocumentIngestionError("active_content", "Macro-enabled Office documents are not accepted")
    except zipfile.BadZipFile as error:
        raise DocumentIngestionError("signature_mismatch", "The file is not a valid DOCX document") from error


def extract_document(
    *,
    content: bytes,
    validated: ValidatedUpload,
    max_pages: int,
) -> ExtractionResult:
    if validated.extension in {".txt", ".md"}:
        return _extract_text(content, validated.extension)
    if validated.extension == ".docx":
        return _extract_docx(content)
    return _extract_pdf(content, max_pages=max_pages)


def _extract_text(content: bytes, extension: str) -> ExtractionResult:
    try:
        value = content.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise DocumentIngestionError("encoding_error", "Text and Markdown files must use UTF-8 encoding") from error
    if len(value.strip()) < 20:
        raise DocumentIngestionError("insufficient_text", "The file does not contain enough readable text")
    return ExtractionResult(
        text=value.strip(),
        method="markdown" if extension == ".md" else "plain_text",
        page_count=None,
        warnings=(),
    )


def _extract_docx(content: bytes) -> ExtractionResult:
    try:
        document = Document(io.BytesIO(content))
    except Exception as error:
        raise DocumentIngestionError("docx_parse_failed", "The DOCX document could not be parsed") from error
    lines: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue
        paragraph_style = paragraph.style
        style = str(paragraph_style.name if paragraph_style else "").lower()
        match = re.match(r"heading\s+(\d+)", style)
        lines.append(f"{'#' * min(int(match.group(1)), 6)} {value}" if match else value)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    extracted = "\n\n".join(lines).strip()
    if len(extracted) < 20:
        raise DocumentIngestionError("insufficient_text", "The DOCX document does not contain enough readable text")
    return ExtractionResult(text=extracted, method="docx_structure", page_count=None, warnings=())


def _extract_pdf(content: bytes, *, max_pages: int) -> ExtractionResult:
    try:
        reader = PdfReader(io.BytesIO(content), strict=True)
        if reader.is_encrypted:
            raise DocumentIngestionError("encrypted_pdf", "Password-protected PDF files are not accepted")
        page_count = len(reader.pages)
        if page_count > max_pages:
            raise DocumentIngestionError("page_limit", f"The PDF exceeds the {max_pages}-page limit")
        extracted_pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except DocumentIngestionError:
        raise
    except Exception as error:
        raise DocumentIngestionError("pdf_parse_failed", "The PDF document could not be parsed safely") from error
    extracted = "\n\n".join(value for value in extracted_pages if value).strip()
    if len(extracted) >= max(80, page_count * 30):
        return ExtractionResult(text=extracted, method="pdf_text", page_count=page_count, warnings=())
    if page_count > 40:
        raise DocumentIngestionError("ocr_page_limit", "Scanned PDFs are limited to 40 pages per upload")
    ocr_text = _ocr_pdf(content, page_count)
    if len(ocr_text) < 20:
        raise DocumentIngestionError("ocr_no_text", "Arabic OCR could not find enough readable text")
    return ExtractionResult(
        text=ocr_text,
        method="pdf_ocr_ara_eng",
        page_count=page_count,
        warnings=("ocr_review_recommended",),
    )


def _ocr_pdf(content: bytes, page_count: int) -> str:
    try:
        document = pdfium.PdfDocument(content)
        values: list[str] = []
        for page_number in range(page_count):
            page = document[page_number]
            bitmap = page.render(scale=2.0)
            image = bitmap.to_pil().convert("RGB")
            values.append(pytesseract.image_to_string(image, lang="ara+eng", config="--psm 6").strip())
            image.close()
            bitmap.close()
            page.close()
        document.close()
        return "\n\n".join(value for value in values if value).strip()
    except Exception as error:
        raise DocumentIngestionError("ocr_failed", "Arabic OCR processing failed") from error


def hash_embedding(value: str) -> tuple[float, ...]:
    normalized = normalize_for_search(value)
    compact = re.sub(r"\s+", " ", normalized).strip()
    terms = compact.split()
    features = terms + [compact[index : index + 4] for index in range(max(0, len(compact) - 3))]
    vector = [0.0] * EMBEDDING_DIMENSION
    for feature in features:
        digest = hashlib.blake2b(feature.encode("utf-8"), digest_size=8).digest()
        index = int.from_bytes(digest[:4], "big") % EMBEDDING_DIMENSION
        sign = 1.0 if digest[4] & 1 else -1.0
        vector[index] += sign
    magnitude = math.sqrt(sum(component * component for component in vector))
    if magnitude == 0:
        return tuple(vector)
    return tuple(round(component / magnitude, 8) for component in vector)


def vector_literal(vector: tuple[float, ...]) -> str:
    if len(vector) != EMBEDDING_DIMENSION:
        raise ValueError("Embedding dimension must be 1024")
    return "[" + ",".join(str(component) for component in vector) + "]"


async def embed_texts(values: list[str], settings: Settings) -> EmbeddingBatch:
    if settings.rag_encoder_mode == "remote":
        try:
            async with httpx.AsyncClient(timeout=settings.rag_encoder_timeout_seconds) as client:
                response = await client.post(f"{settings.encoders_url.rstrip('/')}/embed", json={"texts": values})
                response.raise_for_status()
                payload = response.json()
            vectors = tuple(tuple(float(component) for component in vector) for vector in payload["embeddings"])
            if len(vectors) != len(values) or any(len(vector) != EMBEDDING_DIMENSION for vector in vectors):
                raise ValueError("Encoder returned an invalid embedding shape")
            return EmbeddingBatch(vectors=vectors, provider=str(payload.get("model", "remote-encoder")))
        except (httpx.HTTPError, KeyError, TypeError, ValueError):
            pass
    return EmbeddingBatch(
        vectors=tuple(hash_embedding(value) for value in values),
        provider="local-hash-fallback-v1",
    )


async def rerank_texts(query: str, values: list[str], settings: Settings) -> tuple[tuple[float, ...], str] | None:
    if settings.rag_encoder_mode != "remote" or not values:
        return None
    try:
        async with httpx.AsyncClient(timeout=settings.rag_encoder_timeout_seconds) as client:
            response = await client.post(
                f"{settings.encoders_url.rstrip('/')}/rerank",
                json={"query": query, "documents": values},
            )
            response.raise_for_status()
            payload = response.json()
        scores = tuple(float(score) for score in payload["scores"])
        if len(scores) != len(values):
            raise ValueError("Reranker returned an invalid score count")
        return scores, str(payload.get("model", "remote-reranker"))
    except (httpx.HTTPError, KeyError, TypeError, ValueError):
        return None


def storage_backend(settings: Settings) -> str:
    if settings.object_store_endpoint and settings.object_store_access_key and settings.object_store_secret_key:
        return "s3-compatible"
    return "local-prototype"


async def put_object(
    settings: Settings,
    *,
    object_key: str,
    content: bytes,
    media_type: str,
) -> str:
    backend = storage_backend(settings)
    if backend == "s3-compatible":
        await asyncio.to_thread(_put_minio, settings, object_key, content, media_type)
    else:
        await asyncio.to_thread(_put_local, settings, object_key, content)
    return backend


async def get_object(settings: Settings, *, object_key: str) -> bytes:
    if storage_backend(settings) == "s3-compatible":
        return await asyncio.to_thread(_get_minio, settings, object_key)
    return await asyncio.to_thread(_get_local, settings, object_key)


def _minio_client(settings: Settings) -> Minio:
    parsed = urlsplit(settings.object_store_endpoint)
    endpoint = parsed.netloc or parsed.path
    return Minio(
        endpoint,
        access_key=settings.object_store_access_key,
        secret_key=settings.object_store_secret_key,
        secure=parsed.scheme == "https",
    )


def _put_minio(settings: Settings, object_key: str, content: bytes, media_type: str) -> None:
    client = _minio_client(settings)
    if not client.bucket_exists(settings.object_store_bucket):
        client.make_bucket(settings.object_store_bucket)
    client.put_object(
        settings.object_store_bucket,
        object_key,
        io.BytesIO(content),
        len(content),
        content_type=media_type,
    )


def _get_minio(settings: Settings, object_key: str) -> bytes:
    response = _minio_client(settings).get_object(settings.object_store_bucket, object_key)
    try:
        return bytes(response.read())
    finally:
        response.close()
        response.release_conn()


def _local_path(settings: Settings, object_key: str) -> Path:
    root = Path(settings.object_store_local_root).resolve()
    destination = (root / object_key).resolve()
    if not destination.is_relative_to(root):
        raise DocumentIngestionError("invalid_object_key", "The object storage key is invalid")
    return destination


def _put_local(settings: Settings, object_key: str, content: bytes) -> None:
    destination = _local_path(settings, object_key)
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_bytes(content)


def _get_local(settings: Settings, object_key: str) -> bytes:
    destination = _local_path(settings, object_key)
    if not destination.is_file():
        raise DocumentIngestionError("object_missing", "The original uploaded file is unavailable")
    return destination.read_bytes()
