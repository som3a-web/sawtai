import io
import math

import pytest
from docx import Document

from sawtai.rag.ingestion import (
    DocumentIngestionError,
    extract_document,
    hash_embedding,
    validate_upload,
)
from sawtai.rag.service import structure_chunks


def test_structure_chunks_preserves_headings_and_size_limit() -> None:
    content = "# السياسة العامة\n\n" + " ".join(f"كلمة{i}" for i in range(190)) + "\n\n## الاستثناءات\n\nنص الاستثناء"

    chunks = structure_chunks(content)

    assert len(chunks) == 3
    assert chunks[0][0] == "السياسة العامة"
    assert len(chunks[0][1].split()) == 180
    assert chunks[1][0] == "السياسة العامة"
    assert chunks[2] == ("الاستثناءات", "نص الاستثناء")


def test_upload_validation_rejects_active_and_malicious_content() -> None:
    with pytest.raises(DocumentIngestionError, match="active or embedded"):
        validate_upload(
            filename="policy.pdf",
            content=b"%PDF-1.7\n/JavaScript malicious",
            declared_media_type="application/pdf",
            max_bytes=1_000_000,
        )
    with pytest.raises(DocumentIngestionError, match="malware"):
        validate_upload(
            filename="policy.txt",
            content=b"EICAR-STANDARD-ANTIVIRUS-TEST-FILE",
            declared_media_type="text/plain",
            max_bytes=1_000_000,
        )


def test_docx_structure_extraction_preserves_headings() -> None:
    source = Document()
    source.add_heading("سياسة خدمة المتعاملين", level=1)
    source.add_paragraph("تتم متابعة طلبات المتعاملين خلال يوم عمل واحد وفق الإجراء المعتمد.")
    buffer = io.BytesIO()
    source.save(buffer)
    content = buffer.getvalue()
    validated = validate_upload(
        filename="service-policy.docx",
        content=content,
        declared_media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        max_bytes=1_000_000,
    )

    result = extract_document(content=content, validated=validated, max_pages=20)

    assert result.method == "docx_structure"
    assert result.text.startswith("# سياسة خدمة المتعاملين")
    assert "يوم عمل واحد" in result.text


def test_hash_embedding_is_normalized_and_deterministic() -> None:
    first = hash_embedding("متابعة طلب المتعامل خلال يوم عمل")
    second = hash_embedding("متابعة طلب المتعامل خلال يوم عمل")

    assert first == second
    assert len(first) == 1024
    assert math.isclose(math.sqrt(sum(value * value for value in first)), 1.0, rel_tol=1e-6)
